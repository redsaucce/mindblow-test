import io
import zipfile

from docx import Document as DocxDocument

from app.models.quiz_data import Quiz
from app.models.quiz_question import QuizQuestion

QUIZ_TYPE_LABELS = {
    "multiple_choice": "Multiple Choice",
    "true_false": "True or False",
    "identification": "Identification",
}


def export_quiz(quiz: Quiz, questions: list[QuizQuestion]) -> bytes:
    doc = DocxDocument()

    doc.add_heading(quiz.title, level=1)
    doc.add_paragraph(QUIZ_TYPE_LABELS.get(quiz.quiz_type, quiz.quiz_type))

    if quiz.direction:
        doc.add_paragraph(quiz.direction)

    doc.add_heading("Questions", level=2)
    for q in sorted(questions, key=lambda x: x.order):
        doc.add_paragraph(f"{q.order}. {q.question_text}")
        if quiz.quiz_type == "multiple_choice" and q.options:
            for i, option in enumerate(q.options):
                letter = chr(ord("A") + i)
                doc.add_paragraph(f"    {letter}. {option}")
        elif quiz.quiz_type == "true_false":
            doc.add_paragraph("    A. True")
            doc.add_paragraph("    B. False")

    doc.add_heading("Answer Key", level=2)

    sorted_questions = sorted(questions, key=lambda x: x.order)
    total = len(sorted_questions)
    half = (total + 1) // 2  # left column gets the extra one if odd

    left_col = sorted_questions[:half]
    right_col = sorted_questions[half:]

    table = doc.add_table(rows=half, cols=2)
    table.style = "Table Grid"

    for i in range(half):
        left_q = left_col[i]
        table.rows[i].cells[0].text = f"{left_q.order}. {left_q.correct_answer or ''}"

        if i < len(right_col):
            right_q = right_col[i]
            table.rows[i].cells[1].text = f"{right_q.order}. {right_q.correct_answer or ''}"

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_quizzes_zip(
    quizzes_with_questions: list[tuple[Quiz, list[QuizQuestion]]],
) -> tuple[bytes, list[str]]:
    zip_buffer = io.BytesIO()
    failed_titles: list[str] = []

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for quiz, questions in quizzes_with_questions:
            try:
                docx_bytes = export_quiz(quiz, questions)
                zf.writestr(f"{quiz.title}.docx", docx_bytes)
            except Exception:
                failed_titles.append(quiz.title)

    return zip_buffer.getvalue(), failed_titles